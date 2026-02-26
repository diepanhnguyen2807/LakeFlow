import docx
import re
from pathlib import Path
from typing import Dict, Any, List

from lakeflow.common.jsonio import write_json
from lakeflow.pipelines.processing.chunking import chunk_text


# ==========================================================
# TEXT NORMALIZATION
# ==========================================================

def normalize_text(text: str) -> str:
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Fix chữ dính số (tuyển sinh3 → tuyển sinh 3)
    text = re.sub(r'([A-Za-zÀ-ỹ])(\d)', r'\1 \2', text)

    return text.strip()


# ==========================================================
# TABLE → TEXT (RAG SAFE FORMAT)
# ==========================================================

def table_to_text(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            rows.append(" | ".join(cells))

    if not rows:
        return ""

    return "\n".join(rows)


# ==========================================================
# EXTRACT BLOCKS IN ORIGINAL ORDER
# ==========================================================

def iter_block_items(parent):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# ==========================================================
# MAIN PIPELINE
# ==========================================================

def run_word_pipeline(
    file_hash: str,
    raw_file_path: Path,
    output_dir: Path,
    validation: Dict[str, Any],
) -> None:

    doc = docx.Document(raw_file_path)
    output_dir.mkdir(parents=True, exist_ok=True)

    sections = []
    current_section = None
    section_counter = 0

    # ------------------------------------------------------
    # 1️⃣ Extract structured blocks (paragraph + table)
    # ------------------------------------------------------

    for block in iter_block_items(doc):

        # -------- PARAGRAPH --------
        if isinstance(block, docx.text.paragraph.Paragraph):
            text = block.text.strip()
            if not text:
                continue

            style = block.style.name if block.style else ""

            # Detect Heading
            if style.startswith("Heading"):
                section_counter += 1
                current_section = {
                    "section_id": f"section_{section_counter}",
                    "title": text,
                    "level": int(style.replace("Heading", "").strip() or 1),
                    "content": []
                }
                sections.append(current_section)
            else:
                if current_section is None:
                    section_counter += 1
                    current_section = {
                        "section_id": f"section_{section_counter}",
                        "title": "Mở đầu",
                        "level": 1,
                        "content": []
                    }
                    sections.append(current_section)

                current_section["content"].append(text)

        # -------- TABLE --------
        elif isinstance(block, docx.table.Table):
            table_text = table_to_text(block)
            if table_text and current_section:
                current_section["content"].append("\n[BẢNG]\n" + table_text + "\n")

    # ------------------------------------------------------
    # 2️⃣ Build sections.json
    # ------------------------------------------------------

    section_metadata = []
    for sec in sections:
        section_metadata.append({
            "section_id": sec["section_id"],
            "title": sec["title"],
            "level": sec["level"],
        })

    write_json(output_dir / "sections.json", section_metadata)

    # ------------------------------------------------------
    # 3️⃣ Build chunks.json (Semantic Chunking)
    # ------------------------------------------------------

    final_chunks = []
    chunk_index = 0

    for sec in sections:

        full_text = normalize_text("\n\n".join(sec["content"]))

        chunks = chunk_text(
            full_text,
            chunk_size=600,
            chunk_overlap=100
        )

        for chunk in chunks:
            chunk_index += 1
            final_chunks.append({
                "chunk_id": f"{file_hash}_c{chunk_index}",
                "text": chunk,
                "section_id": sec["section_id"],
                "file_hash": file_hash,
                "token_estimate": len(chunk.split()),
            })

    write_json(output_dir / "chunks.json", final_chunks)

    print(f"[PROCESS][WORD] {raw_file_path.name} → {len(final_chunks)} semantic chunks created.")

# code trước 24/2/2026
# import docx
# import json
# from pathlib import Path
# from typing import Dict, Any, List

# from lakeflow.common.jsonio import write_json
# from lakeflow.pipelines.processing.chunking import chunk_text

# def run_word_pipeline(
#     file_hash: str,
#     raw_file_path: Path,
#     output_dir: Path,
#     validation: Dict[str, Any],
# ) -> None:
#     """
#     Xử lý Word (.docx) → sinh dữ liệu AI-ready (300_processed).
#     Phân tách văn bản và bảng biểu theo cấu trúc LakeFlow tiêu chuẩn.
#     """

#     # ---------- 1. Load Word Document ----------
#     doc = docx.Document(raw_file_path)
#     output_dir.mkdir(parents=True, exist_ok=True)

#     # ---------- 2. Extract Tables (Tương tự Excel Pipeline) ----------
#     all_tables = []
#     for i, table in enumerate(doc.tables):
#         rows_data = []
#         for row in table.rows:
#             rows_data.append([cell.text.strip() for cell in row.cells])
        
#         if rows_data:
#             all_tables.append({
#                 "table_id": f"{file_hash}_t{i+1}",
#                 "title": f"Bảng số {i+1} trong tài liệu Word",
#                 "headers": rows_data[0] if len(rows_data) > 0 else [],
#                 "rows": rows_data[1:] if len(rows_data) > 1 else [],
#                 "source_file": raw_file_path.name
#             })
    
#     write_json(output_dir / "tables.json", all_tables)

#     # ---------- 3. Extract & Clean Text ----------
#     # Lấy text từ các đoạn văn (paragraphs)
#     paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
#     full_text = "\n\n".join(paragraphs)

#     (output_dir / "clean_text.txt").write_text(full_text, encoding="utf-8")

#     # ---------- 4. Build sections.json ----------
#     # Đối với Word, chúng ta có thể coi toàn bộ là một Section chính
#     # Hoặc có thể mở rộng để detect Heading (nếu cần)
#     sections = [
#         {
#             "section_id": "main_content",
#             "title": "Nội dung văn bản chính",
#             "level": 1,
#         }
#     ]
#     write_json(output_dir / "sections.json", sections)

#     # ---------- 5. Build chunks.json (Sử dụng logic chunking chuẩn) ----------
#     # Chia nhỏ văn bản thành các đoạn phù hợp với context window của AI
#     raw_chunks = chunk_text(
#         full_text,
#         chunk_size=600,
#         chunk_overlap=100
#     )

#     final_chunks = []
#     for i, text_segment in enumerate(raw_chunks):
#         final_chunks.append({
#             "chunk_id": f"{file_hash}_c{i+1}",
#             "text": text_segment,
#             "section_id": "main_content",
#             "file_hash": file_hash,
#             "token_estimate": len(text_segment.split()),
#         })

#     write_json(output_dir / "chunks.json", final_chunks)

#     print(f"   [PROCESS][WORD] {raw_file_path.name} -> {len(final_chunks)} chunks created.")
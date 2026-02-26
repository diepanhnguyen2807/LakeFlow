# frontend/streamlit/pages/data_lake_explorer.py

from pathlib import Path
import hashlib
import json
import os
import sqlite3

import pandas as pd
import streamlit as st

from config.settings import DATA_ROOT
from state.session import require_login
from services.api_client import get_data_path_from_api

# File viewer: giới hạn kích thước (tránh treo)
MAX_VIEW_TEXT_BYTES = 10 * 1024 * 1024   # 10 MB cho txt/json/jsonl
MAX_VIEW_NPY_BYTES = 5 * 1024 * 1024     # 5 MB cho npy
MAX_VIEW_PDF_BYTES = 50 * 1024 * 1024    # 50 MB cho pdf
MAX_JSONL_LINES = 500

# =====================================================
# DATA ZONES
# =====================================================

ZONE_NAMES = [
    "000_inbox",
    "100_raw",
    "200_staging",
    "300_processed",
    "400_embeddings",
    "500_catalog",
]

def _zones_from_root(root: Path) -> dict[str, Path]:
    return {z: root / z for z in ZONE_NAMES}

# Fallback khi chưa gọi API (module load)
ZONES = _zones_from_root(DATA_ROOT)

MAX_TREE_DEPTH = 30  # Giới hạn độ sâu tránh đệ quy vô hạn
CACHE_TTL_TREE = 90  # Giây cache cho list_dir (NAS chậm)

# Bước pipeline: 0=inbox→raw, 1=raw→staging, 2=staging→processed, 3=processed→embeddings, 4=embeddings→Qdrant
PIPELINE_STEP_LABELS = {
    0: "Step 0 (Inbox → Raw)",
    1: "Step 1 (Raw → Staging)",
    2: "Step 2 (Staging → Processed)",
    3: "Step 3 (Processed → Embeddings)",
    4: "Step 4 (Embeddings → Qdrant)",
}

# Cột trạng thái từng bước cho file trong 000_inbox
INBOX_STEP_COLUMNS = ["Ingest", "Staging", "Processed", "Embeddings", "Qdrant"]


# =====================================================
# CACHE ĐỌC NAS (giảm lag, tránh đọc lại cùng path)
# =====================================================

@st.cache_data(ttl=CACHE_TTL_TREE)
def _list_dir_cached(path_str: str) -> tuple[list[str], list[tuple[str, int]]]:
    """
    Đọc thư mục từ NAS, trả về (tên thư mục, [(tên file, size)]).
    Dùng thread để không block UI; kết quả được cache.
    """
    path = Path(path_str)
    dirs, files = [], []
    try:
        for p in sorted(path.iterdir()):
            if p.name.startswith("."):
                continue
            if p.is_dir():
                dirs.append(p.name)
            elif p.is_file():
                try:
                    files.append((p.name, p.stat().st_size))
                except OSError:
                    files.append((p.name, 0))
    except (PermissionError, OSError):
        pass
    return (sorted(dirs, key=str.lower), sorted(files, key=lambda x: x[0].lower()))


# =====================================================
# TREE VIEW — LAZY LOADING
# =====================================================

def _format_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} B"
    if size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / (1024 * 1024):.1f} MB"


def _path_to_key(path_str: str) -> str:
    """Tạo key widget duy nhất từ path (tránh trùng khi path dài)."""
    return hashlib.md5(path_str.encode()).hexdigest()[:24]


@st.cache_data(ttl=120)
def _sha256_file_cached(path_str: str) -> str | None:
    """Tính SHA256 của file (để đối chiếu với raw_objects). Cache 120s."""
    path = Path(path_str)
    if not path.is_file():
        return None
    try:
        h = hashlib.sha256()
        with path.open("rb") as f:
            while chunk := f.read(1024 * 1024):
                h.update(chunk)
        return h.hexdigest()
    except (OSError, PermissionError):
        return None


def get_inbox_file_pipeline_steps(file_path: Path, domain: str) -> dict[str, str]:
    """
    Với file trong 000_inbox: trả về từng bước đã xử lý hay chưa.
    Keys: Ingest, Staging, Processed, Embeddings, Qdrant. Value: "✓" hoặc "" (Qdrant có thể "?" nếu không biết).
    """
    result = {k: "" for k in INBOX_STEP_COLUMNS}
    path_str = str(file_path.resolve())
    file_hash = _sha256_file_cached(path_str)
    if not file_hash:
        return result
    root = DATA_ROOT
    catalog_db = root / "500_catalog" / "catalog.sqlite"
    # Step 0: Ingest (đã có trong raw_objects)
    try:
        if catalog_db.exists():
            conn = sqlite3.connect(f"file:{catalog_db}?mode=ro", uri=True, timeout=5)
            cur = conn.execute("SELECT 1 FROM raw_objects WHERE hash = ? LIMIT 1", (file_hash,))
            if cur.fetchone():
                result["Ingest"] = "✓"
            conn.close()
    except Exception:
        pass
    if result["Ingest"] != "✓":
        return result
    # Steps 1–3: kiểm tra thư mục; hỗ trợ cả domain/hash và hash (cấu trúc cũ)
    _staging_dir = root / "200_staging" / domain / file_hash if domain and domain != "." else root / "200_staging" / file_hash
    _staging_alt = root / "200_staging" / file_hash if domain and domain != "." else None
    if (_staging_dir / "validation.json").exists() or (_staging_alt and (_staging_alt / "validation.json").exists()):
        result["Staging"] = "✓"
    _processed_dir = root / "300_processed" / domain / file_hash if domain and domain != "." else root / "300_processed" / file_hash
    _processed_alt = root / "300_processed" / file_hash if domain and domain != "." else None
    if (_processed_dir / "chunks.json").exists() or (_processed_alt and (_processed_alt / "chunks.json").exists()):
        result["Processed"] = "✓"
    _emb_dir = root / "400_embeddings" / domain / file_hash if domain and domain != "." else root / "400_embeddings" / file_hash
    _emb_alt = root / "400_embeddings" / file_hash if domain and domain != "." else None
    if (_emb_dir / "embedding.npy").exists() or (_emb_alt and (_emb_alt / "embedding.npy").exists()):
        result["Embeddings"] = "✓"
    # Step 4: Qdrant — không có trong catalog, để trống hoặc "?"
    result["Qdrant"] = "?" if result["Embeddings"] == "✓" else ""
    return result


def get_raw_file_pipeline_steps(file_path: Path, domain: str) -> dict[str, str]:
    """
    Với file trong 100_raw: Ingest luôn ✓ (đã ở Raw), các bước sau kiểm tra theo domain/hash.
    file_hash = file_path.stem (tên file không extension).
    """
    result = {k: "" for k in INBOX_STEP_COLUMNS}
    result["Ingest"] = "✓"  # Đã ingest (file đang ở Raw)
    file_hash = file_path.stem
    domain = domain or "."
    root = DATA_ROOT
    # Hỗ trợ cả domain/hash và hash (cấu trúc cũ); backend ghi embedding.npy (không có s)
    _staging = root / "200_staging" / domain / file_hash if domain != "." else root / "200_staging" / file_hash
    _staging_alt = root / "200_staging" / file_hash if domain != "." else None
    if (_staging / "validation.json").exists() or (_staging_alt and (_staging_alt / "validation.json").exists()):
        result["Staging"] = "✓"
    _processed = root / "300_processed" / domain / file_hash if domain != "." else root / "300_processed" / file_hash
    _processed_alt = root / "300_processed" / file_hash if domain != "." else None
    if (_processed / "chunks.json").exists() or (_processed_alt and (_processed_alt / "chunks.json").exists()):
        result["Processed"] = "✓"
    _emb = root / "400_embeddings" / domain / file_hash if domain != "." else root / "400_embeddings" / file_hash
    _emb_alt = root / "400_embeddings" / file_hash if domain != "." else None
    if (_emb / "embedding.npy").exists() or (_emb_alt and (_emb_alt / "embedding.npy").exists()):
        result["Embeddings"] = "✓"
    result["Qdrant"] = "?" if result["Embeddings"] == "✓" else ""
    return result


def render_folder_tree(
    root: Path,
    zone_name: str,
    expanded_set: set[str],
    current_folder: str | None,
    zone_root: Path,
    depth: int = 0,
) -> None:
    """
    Chỉ hiển thị cây thư mục (không có file). Bấm thư mục = mở/đóng + chọn để hiển thị danh sách file bên phải.
    """
    if depth >= MAX_TREE_DEPTH:
        st.caption("… (đạt giới hạn độ sâu)")
        return

    path_str = str(root.resolve())
    try:
        dir_names, _ = _list_dir_cached(path_str)
    except Exception as e:
        st.caption(f"⚠️ Lỗi: {e}")
        return

    indent = "  " * depth  # 2 space cho gọn

    for d_name in dir_names:
        child_path = root / d_name
        child_str = str(child_path.resolve())
        key_suffix = _path_to_key(child_str)
        is_selected = current_folder == child_str
        icon = "▼" if child_str in expanded_set else "▶"
        try:
            child_dirs, child_files = _list_dir_cached(child_str)
            count = len(child_dirs) + len(child_files)
            count_str = f" ({count})"
        except Exception:
            count_str = ""
        label = f"{indent}{icon} {d_name}{count_str}" + (" ✓" if is_selected else "")

        if st.button(label, key=f"tree_{zone_name}_{key_suffix}", type="primary" if is_selected else "secondary"):
            if child_str in expanded_set:
                expanded_set.discard(child_str)
            else:
                expanded_set.add(child_str)
            st.session_state["datalake_current_folder"] = child_str
            st.rerun()

        if child_str in expanded_set:
            render_folder_tree(child_path, zone_name, expanded_set, current_folder, zone_root, depth + 1)


def render_file_list(
    folder_path: Path,
    zone_name: str,
    zone_root: Path,
    selected_file: str | None,
) -> None:
    """
    Hiển thị danh sách file trong thư mục dạng bảng; chọn dòng để xem nội dung.
    """
    if not _is_safe_path(folder_path, zone_root):
        st.warning("Thư mục không thuộc zone.")
        return
    try:
        _, file_infos = _list_dir_cached(str(folder_path.resolve()))
    except Exception as e:
        st.warning(f"Không đọc được thư mục: {e}")
        return

    if not file_infos:
        st.info("Thư mục trống hoặc không có file.")
        return

    # 000_inbox hoặc 100_raw (Gốc hoặc thư mục domain): bảng có cột từng bước với ✓
    # 100_raw bỏ qua bước Ingest (luôn ✓ vì file đã ở Raw)
    # Zone khác: bảng có cột Bước pipeline (mô tả text)
    rows = []
    is_inbox_zone = zone_name == "000_inbox"
    is_raw_zone = zone_name == "100_raw"
    use_step_columns = is_inbox_zone or is_raw_zone
    # Domain = segment đầu tiên dưới zone (để đúng cả khi xem file trong thư mục con, vd. 000_inbox/education/2024/)
    try:
        rel = folder_path.resolve().relative_to(zone_root.resolve())
        domain = rel.parts[0] if rel.parts else "."
    except (ValueError, OSError):
        domain = "." if folder_path.resolve() == zone_root.resolve() else folder_path.name
    for f_name, size in file_infos:
        file_path = folder_path / f_name
        _path = str(file_path.resolve())
        if use_step_columns and is_inbox_zone:
            steps = get_inbox_file_pipeline_steps(file_path, domain)
            rows.append({
                "Tên file": f_name,
                "Kích thước": _format_size(size),
                **{col: steps.get(col, "") for col in INBOX_STEP_COLUMNS},
                "_path": _path,
            })
        elif use_step_columns and is_raw_zone:
            steps = get_raw_file_pipeline_steps(file_path, domain)
            rows.append({
                "Tên file": f_name,
                "Kích thước": _format_size(size),
                **{col: steps.get(col, "") for col in INBOX_STEP_COLUMNS},
                "_path": _path,
            })
        else:
            step_label = get_pipeline_step_for_path(
                file_path, zone_name, data_root=zone_root.parent
            )
            step_str = step_label if step_label else "—"
            rows.append({
                "Tên file": f_name,
                "Kích thước": _format_size(size),
                "Bước pipeline": step_str,
                "_path": _path,
            })

    if use_step_columns:
        df = pd.DataFrame([{k: r[k] for k in ["Tên file", "Kích thước"] + INBOX_STEP_COLUMNS} for r in rows])
        if is_inbox_zone:
            st.caption("✓ = đã xử lý bước đó. Ingest = Step 0 (→ Raw), Staging = Step 1, Processed = Step 2, Embeddings = Step 3, Qdrant = Step 4 (?) nếu chưa xác nhận.")
        else:
            st.caption("✓ = đã xử lý bước đó. File trong Raw nên Ingest luôn ✓. Staging = Step 1, Processed = Step 2, Embeddings = Step 3, Qdrant = Step 4 (?) nếu chưa xác nhận.")
    else:
        df = pd.DataFrame([{"Tên file": r["Tên file"], "Kích thước": r["Kích thước"], "Bước pipeline": r["Bước pipeline"]} for r in rows])

    st.markdown(f"**Files trong** `{folder_path.name}`")
    st.dataframe(df, use_container_width=True, hide_index=True)

    chosen = st.selectbox(
        "Chọn file để xem nội dung",
        options=[r["_path"] for r in rows],
        format_func=lambda p: Path(p).name,
        index=next((i for i, r in enumerate(rows) if r["_path"] == selected_file), 0),
        key=f"file_sel_{zone_name}_{_path_to_key(str(folder_path.resolve()))}",
    )
    if chosen:
        st.session_state["datalake_selected_file"] = chosen


def get_pipeline_step_for_path(
    file_path: Path, zone_name: str, data_root: Path | None = None
) -> str | None:
    """
    Xác định file đã được xử lý tới bước nào trong pipeline (0–4) bằng catalog và kiểm tra file hệ thống.
    Trả về mô tả bước hoặc None nếu không xác định được.
    data_root: nếu có thì dùng (path từ backend); không thì dùng DATA_ROOT từ env.
    """
    root = data_root if data_root is not None else DATA_ROOT
    staging_root = root / "200_staging"
    processed_root = root / "300_processed"
    embeddings_root = root / "400_embeddings"
    catalog_db = root / "500_catalog" / "catalog.sqlite"

    if zone_name == "000_inbox":
        return "Chưa chạy Step 0 (file vẫn trong Inbox)"
    if zone_name not in ("100_raw", "200_staging", "300_processed", "400_embeddings"):
        return None

    file_hash = None
    domain = None
    zone_root = root / zone_name if data_root is not None else ZONES.get(zone_name)
    if not zone_root:
        return None
    try:
        rel = file_path.resolve().relative_to(zone_root.resolve())
        parts = rel.parts
        if zone_name == "100_raw" and file_path.is_file():
            file_hash = file_path.stem
            domain = parts[0] if len(parts) >= 2 else None
        elif zone_name in ("200_staging", "300_processed", "400_embeddings"):
            if file_path.is_dir():
                file_hash = file_path.name
                domain = parts[0] if len(parts) >= 2 else None
            else:
                file_hash = file_path.parent.name
                domain = parts[0] if len(parts) >= 2 else None
    except ValueError:
        return None
    if not file_hash:
        return None

    domain = domain or "."
    step = -1
    if not catalog_db.exists():
        return "Catalog chưa có (chưa chạy Step 0)"

    try:
        conn = sqlite3.connect(f"file:{catalog_db}?mode=ro", uri=True, timeout=5)
        cur = conn.execute("SELECT 1 FROM raw_objects WHERE hash = ? LIMIT 1", (file_hash,))
        if cur.fetchone():
            step = 0
        conn.close()
    except Exception:
        return "Không đọc được Catalog"

    if step < 0:
        return "Chưa có trong Catalog (chưa chạy Step 0)"

    domain = domain or "."
    # Hỗ trợ cả domain/hash và hash (cấu trúc cũ); backend ghi embedding.npy (không có s)
    _staging = staging_root / domain / file_hash if domain != "." else staging_root / file_hash
    _staging_alt = staging_root / file_hash if domain != "." else None
    if (_staging / "validation.json").exists() or (_staging_alt and (_staging_alt / "validation.json").exists()):
        step = 1
    _processed = processed_root / domain / file_hash if domain != "." else processed_root / file_hash
    _processed_alt = processed_root / file_hash if domain != "." else None
    if (_processed / "chunks.json").exists() or (_processed_alt and (_processed_alt / "chunks.json").exists()):
        step = 2
    _emb = embeddings_root / domain / file_hash if domain != "." else embeddings_root / file_hash
    _emb_alt = embeddings_root / file_hash if domain != "." else None
    if (_emb / "embedding.npy").exists() or (_emb_alt and (_emb_alt / "embedding.npy").exists()):
        step = 3
    # Step 4: cần query Qdrant để xác nhận — hiện không có trong catalog

    return PIPELINE_STEP_LABELS.get(step, f"Step {step}")


def _is_safe_path(file_path: Path, zone_root: Path) -> bool:
    """Đảm bảo file nằm trong zone (tránh path traversal)."""
    try:
        return file_path.resolve().is_relative_to(zone_root.resolve())
    except (ValueError, OSError):
        return False


def render_file_content(file_path: Path) -> None:
    """
    Hiển thị nội dung file theo định dạng: txt, json, jsonl, npy, pdf, csv.
    Giới hạn kích thước để tránh treo.
    """
    if not file_path.is_file():
        st.warning("File không tồn tại hoặc không đọc được.")
        return

    try:
        size = file_path.stat().st_size
    except OSError:
        st.error("Không đọc được thông tin file.")
        return

    suffix = file_path.suffix.lower()

    # ---------- TXT ----------
    if suffix == ".txt":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File quá lớn ({size / (1024*1024):.1f} MB). Chỉ hỗ trợ xem file ≤ {MAX_VIEW_TEXT_BYTES // (1024*1024)} MB.")
            _download_button(file_path)
            return
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            st.code(text, language="text")
        except Exception as e:
            st.error(f"Lỗi đọc file: {e}")

    # ---------- JSON ----------
    elif suffix == ".json":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File quá lớn. Chỉ hỗ trợ xem file ≤ {MAX_VIEW_TEXT_BYTES // (1024*1024)} MB.")
            _download_button(file_path)
            return
        try:
            with file_path.open("r", encoding="utf-8") as f:
                data = json.load(f)
            st.json(data)
        except Exception as e:
            st.error(f"Lỗi đọc JSON: {e}")

    # ---------- JSONL ----------
    elif suffix == ".jsonl":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File quá lớn. Chỉ hiển thị tối đa {MAX_JSONL_LINES} dòng đầu.")
        try:
            lines = []
            with file_path.open("r", encoding="utf-8", errors="replace") as f:
                for i, line in enumerate(f):
                    if i >= MAX_JSONL_LINES:
                        st.caption(f"… Chỉ hiển thị {MAX_JSONL_LINES} dòng đầu. Tổng file có thể nhiều hơn.")
                        break
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        lines.append(json.loads(line))
                    except json.JSONDecodeError:
                        lines.append({"raw": line})
            if lines:
                st.dataframe(lines, use_container_width=True)
            else:
                st.info("File rỗng hoặc không có dòng JSON hợp lệ.")
        except Exception as e:
            st.error(f"Lỗi đọc file: {e}")

    # ---------- NPY ----------
    elif suffix == ".npy":
        if size > MAX_VIEW_NPY_BYTES:
            st.warning(f"File quá lớn ({size / (1024*1024):.1f} MB). Chỉ hỗ trợ xem file ≤ {MAX_VIEW_NPY_BYTES // (1024*1024)} MB.")
            _download_button(file_path)
            return
        try:
            import numpy as np
            arr = np.load(file_path, allow_pickle=False)
            st.write("**Shape:**", arr.shape)
            st.write("**Dtype:**", str(arr.dtype))
            if arr.size <= 100:
                st.write("**Dữ liệu:**")
                st.write(arr)
            else:
                st.write("**Mẫu (100 phần tử đầu):**")
                st.write(arr.flat[:100])
        except ImportError:
            st.info("Cần cài `numpy` để xem file .npy. Bạn có thể tải file xuống.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Lỗi đọc file .npy: {e}")

    # ---------- PDF ----------
    elif suffix == ".pdf":
        if size > MAX_VIEW_PDF_BYTES:
            st.warning(f"File quá lớn. Chỉ hỗ trợ thông tin file ≤ {MAX_VIEW_PDF_BYTES // (1024*1024)} MB.")
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            n_pages = len(reader.pages)
            st.write(f"**Số trang:** {n_pages}")
            _download_button(file_path)
        except ImportError:
            st.info("Cần cài `pypdf` để xem thông tin PDF. Bạn có thể tải file xuống.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Lỗi đọc PDF: {e}")
            _download_button(file_path)

    # ---------- CSV ----------
    elif suffix == ".csv":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File quá lớn. Chỉ hiển thị một phần.")
        try:
            import pandas as pd
            df = pd.read_csv(file_path, nrows=1000, encoding="utf-8", on_bad_lines="skip")
            st.dataframe(df, use_container_width=True)
            if size > 1024 * 1024:
                st.caption("Chỉ hiển thị 1000 dòng đầu.")
        except ImportError:
            st.info("Cần cài `pandas` để xem CSV. Bạn có thể tải file xuống.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Lỗi đọc CSV: {e}")
    # Cao Hương thêm
    elif suffix == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
            # Trích xuất toàn bộ văn bản từ các đoạn (paragraphs)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Trích xuất thêm từ bảng (để không sót thông tin như số nguyện vọng)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    full_text.append(f" [Bảng] {row_text}")
            
            st.write(f"**Định dạng:** Microsoft Word")
            st.text_area("Nội dung trích xuất:", value="\n".join(full_text), height=400)
            _download_button(file_path)
        except ImportError:
            st.info("Cần cài `python-docx` để xem nội dung Word. Bạn có thể tải file xuống.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Lỗi đọc Word: {e}")
            _download_button(file_path)
    elif suffix in [".xls", ".xlsx"]:
        try:
            import pandas as pd
            # Đọc sheet đầu tiên của Excel
            df = pd.read_excel(file_path)
            
            st.write(f"**Định dạng:** Microsoft Excel")
            st.write(f"**Số dòng:** {len(df)} | **Số cột:** {len(df.columns)}")
            
            # Hiển thị bảng dữ liệu trực quan
            st.dataframe(df, use_container_width=True)
            _download_button(file_path)
        except ImportError:
            st.info("Cần cài `pandas`, `openpyxl` và `xlrd` để xem Excel. Bạn có thể tải file xuống.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Lỗi đọc Excel: {e}")
            _download_button(file_path)
    #============================================CaoHuong===hết===
    
    # ---------- Khác ----------
    else:
        st.info(f"Định dạng `{suffix}` chưa hỗ trợ xem trực tiếp. Bạn có thể tải file xuống.")
        _download_button(file_path)


def _download_button(file_path: Path) -> None:
    try:
        data = file_path.read_bytes()
        st.download_button(
            "⬇️ Tải file xuống",
            data=data,
            file_name=file_path.name,
            mime="application/octet-stream",
            key=f"dl_{hashlib.md5(str(file_path).encode()).hexdigest()[:16]}",
        )
    except Exception:
        pass


# =====================================================
# MAIN PAGE
# =====================================================

def render():
    if not require_login():
        return

    # Lấy data root từ backend (đúng path khi chạy dev; tránh /data mặc định)
    if "data_lake_root" not in st.session_state:
        api_path = get_data_path_from_api()
        if api_path:
            st.session_state["data_lake_root"] = Path(api_path).expanduser().resolve()
        else:
            st.session_state["data_lake_root"] = DATA_ROOT
    effective_root = st.session_state["data_lake_root"]
    zones = _zones_from_root(effective_root)

    st.header("🗂️ Data Lake Explorer")
    st.caption(
        "Xem cấu trúc Data Lake. Khi bấm vào file, hiển thị nội dung và **bước pipeline** (từ Catalog: đã chạy Step 0–3). "
        "Catalog (500_catalog) chứa raw_objects, ingest_log."
    )

    # --------------------------------------------------
    # SELECT ZONE
    # --------------------------------------------------
    zone_name = st.selectbox(
        "📂 Chọn data zone",
        list(zones.keys()),
    )

    zone_path = zones[zone_name]

    if not zone_path.exists():
        st.warning(f"Zone chưa tồn tại: {zone_path}")
        return

    st.subheader(f"📁 {zone_name}")

    if "datalake_expanded" not in st.session_state:
        st.session_state.datalake_expanded = {}
    expanded_set = st.session_state.datalake_expanded.setdefault(zone_name, set())
    zone_root_str = str(zone_path.resolve())
    current_folder = st.session_state.get("datalake_current_folder")
    if not current_folder:
        current_folder = zone_root_str
    else:
        try:
            Path(current_folder).resolve().relative_to(zone_path.resolve())
        except (ValueError, OSError, TypeError):
            current_folder = zone_root_str
    if not current_folder:
        current_folder = zone_root_str

    # --------------------------------------------------
    # LAYOUT 2 CỘT: Cây thư mục (gọn) trái | Bảng file + nội dung phải
    # --------------------------------------------------
    col_tree, col_files = st.columns([0.9, 2.1])

    with col_tree:
        st.markdown("**📁 Thư mục**")
        try:
            root_dirs, root_files = _list_dir_cached(zone_root_str)
            root_count = len(root_dirs) + len(root_files)
            root_label = f"📂 Gốc ({root_count})"
        except Exception:
            root_label = "📂 Gốc"
        if st.button(root_label, key="datalake_root", help="Xem file trong thư mục gốc zone"):
            st.session_state["datalake_current_folder"] = zone_root_str
            st.rerun()
        st.divider()
        with st.spinner("Đang tải cây..."):
            render_folder_tree(zone_path, zone_name, expanded_set, current_folder, zone_path)

    with col_files:
        st.markdown("**📄 Danh sách file**")
        folder_path = Path(current_folder)

        with st.spinner("Đang tải danh sách..."):
            render_file_list(folder_path, zone_name, zone_path, st.session_state.get("datalake_selected_file"))

        # Nội dung file khi đã chọn
        selected = st.session_state.get("datalake_selected_file")
        if selected:
            sel_path = Path(selected)
            if sel_path.is_file() and _is_safe_path(sel_path, zone_path):
                st.divider()
                st.subheader(f"📄 `{sel_path.name}`")
                step_info = get_pipeline_step_for_path(
                    sel_path, zone_name, data_root=effective_root
                )
                if step_info:
                    st.caption(f"🔄 **Pipeline:** {step_info}")
                if st.button("✕ Đóng xem file", key="datalake_close_file"):
                    del st.session_state["datalake_selected_file"]
                    st.rerun()
                render_file_content(sel_path)

    if zone_name == "500_catalog":
        st.caption("Xem nội dung SQLite (catalog) tại **🗄️ SQLite Viewer** trên thanh bên trái.")

from pathlib import Path
from typing import Dict, Any
import docx
from docx.opc.exceptions import PackageNotFoundError

class StagingError(RuntimeError):
    """Lỗi staging với lý do rõ ràng, khớp với cách bắt lỗi của project."""

def analyze_word(file_path: Path) -> Dict[str, Any]:
    """
    Phân tích cấu trúc file Word chi tiết (200_staging).
    Bao gồm kiểm tra lỗi định dạng, trích xuất metadata và đánh giá nội dung.
    """
    if not file_path.exists():
        raise StagingError(f"File Word không tồn tại: {file_path}")

    try:
        doc = docx.Document(file_path)
    except PackageNotFoundError:
        raise StagingError(f"File không đúng định dạng Word (.docx) hoặc bị hỏng cấu trúc ZIP.")
    except Exception as e:
        raise StagingError(f"Lỗi phân tích Word: {str(e)}")

    # 1. Thu thập Metadata (Tương tự metadata của PDF)
    prop = doc.core_properties
    metadata = {
        "author": prop.author,
        "created": str(prop.created) if prop.created else None,
        "last_modified_by": prop.last_modified_by,
        "revision": prop.revision,
        "title": prop.title,
    }

    # 2. Phân tích cấu trúc nội dung
    paragraphs = doc.paragraphs
    tables = doc.tables
    
    # Kiểm tra xem có thực sự có chữ không (đề phòng file trắng)
    text_content = [p.text for p in paragraphs if p.text.strip()]
    has_text = len(text_content) > 0
    
    # Tính toán tỷ lệ (tương tự với file PDF)
    paragraph_count = len(paragraphs)
    table_count = len(tables)

    return {
        "file_type": "docx",
        "status": "valid",
        "metadata": metadata,
        
        # Thống kê số lượng
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "has_text_layer": has_text,
        
        # Quyết định kỹ thuật cho Step 2 (Processing)
        "requires_table_extraction": table_count > 0,
        "requires_text_processing": has_text,
        "is_empty": not (has_text or table_count > 0),
        
        # Thông tin bổ sung cho logs
        "word_version": "Office Open XML (OOXML)"
    }